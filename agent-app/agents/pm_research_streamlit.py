import streamlit as st
import pandas as pd
from phi.agent import Agent
from phi.model.openai import OpenAIChat
from phi.tools.googlesearch import GoogleSearch
from phi.tools.yfinance import YFinanceTools
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import re
import os
import base64
from io import BytesIO

# Initialize agents
def initialize_agents():
    web_agent = Agent(
        name="Web Agent",
        role="Search the web for latest information and news",
        model=OpenAIChat(id="gpt-4o"),
        tools=[GoogleSearch()],
        instructions=[
            "Search for latest news and information about the given topic",
            "Provide 3-5 key findings with dates and sources",
            "Format information in clear, digestible bullet points"
        ],
        show_tool_calls=True,
        markdown=True,
    )

    finance_agent = Agent(
        name="Finance Agent",
        role="Analyze financial data and market trends",
        model=OpenAIChat(id="gpt-4o"),
        tools=[YFinanceTools(stock_price=True, analyst_recommendations=True, company_info=True)],
        instructions=[
            "Analyze financial metrics and market data",
            "Present data in clear tables",
            "Highlight key financial insights and trends"
        ],
        show_tool_calls=True,
        markdown=True,
    )

    tech_market_agent = Agent(
        name="Technology and Market Opportunity Expert",
        role="Analyze technology trends, market dynamics, and identify value creation opportunities",
        model=OpenAIChat(id="gpt-4o"),
        tools=[GoogleSearch()],
        instructions=[
            "Provide a structured market analysis with these specific sections:",
            
            "1. MARKET SIZE & GROWTH",
            "- Current global market size with specific dollar amount",
            "- Year-over-year growth rate (CAGR)",
            "- 5-year market size projection",
            "- Break down by major geographic regions",
            
            "2. MARKET SEGMENTS",
            "- List top 3-5 market segments with size/share",
            "- Identify fastest growing segments",
            "- Key drivers for each segment",
            
            "3. COMPETITIVE LANDSCAPE",
            "- Market share of top 5 players",
            "- Recent funding rounds and valuations",
            "- Key partnerships and acquisitions",
            
            "4. GROWTH DRIVERS & TRENDS",
            "- List specific technological advancements",
            "- Regulatory impacts",
            "- Customer demand patterns"
        ],
        show_tool_calls=True,
        markdown=True,
    )

    value_capture_agent = Agent(
        name="Value Capture Strategist",
        role="Develop strategies for IP protection, market positioning, and competitive advantage",
        model=OpenAIChat(id="gpt-4o"),
        tools=[GoogleSearch()],
        instructions=[
            "Focus on IP protection strategies",
            "Develop market positioning recommendations",
            "Identify competitive advantages",
            "Provide actionable strategic recommendations",
            "Always include sources"
        ],
        show_tool_calls=True,
        markdown=True,
    )

    org_design_agent = Agent(
        name="Organizational Design Architect",
        role="Design optimal organizational structures and collaboration networks",
        model=OpenAIChat(id="gpt-4o"),
        tools=[GoogleSearch()],
        instructions=[
            "Design team structures and collaboration frameworks",
            "Optimize for innovation and value delivery",
            "Consider organizational culture and dynamics",
            "Provide practical implementation steps",
            "Always include sources"
        ],
        show_tool_calls=True,
        markdown=True,
    )
    
    return web_agent, finance_agent, tech_market_agent, value_capture_agent, org_design_agent

class ReportFormatter:
    def __init__(self):
        self.document = Document()
        style = self.document.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
    def clean_output(self, text):
        cleaned = re.sub(r'Running:.*?(?=\n\n)', '', text, flags=re.DOTALL)
        cleaned = re.sub(r'• transfer_task_to.*?\.\.\..*?\n', '', cleaned)
        cleaned = re.sub(r'^\s*\n', '', cleaned)
        return cleaned.strip()
        
    def add_section(self, title, content):
        section_heading = self.document.add_heading(level=1)
        heading_run = section_heading.add_run(title)
        heading_run.bold = True
        
        cleaned_content = self.clean_output(content)
        self.document.add_paragraph(cleaned_content)
        self.document.add_paragraph('_' * 50)

    def format_document(self, title, agent_outputs):
        title_heading = self.document.add_heading(level=0)
        title_run = title_heading.add_run(title)
        title_run.bold = True
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_paragraph = self.document.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_paragraph.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        
        self.document.add_paragraph('_' * 50)
        
        for agent_name, output in agent_outputs:
            self.add_section(f"{agent_name} Analysis", output)

    def save_to_bytes(self):
        docx_bytes = BytesIO()
        self.document.save(docx_bytes)
        docx_bytes.seek(0)
        return docx_bytes

def run_analysis(business_type, progress_bar, status_text):
    web_agent, finance_agent, tech_market_agent, value_capture_agent, org_design_agent = initialize_agents()
    
    agent_outputs = []
    progress_step = 0.2  # Changed from 100/5 to 0.2 (20% per step)
    
    try:
        # Web Agent
        status_text.text("Gathering latest news...")
        web_prompt = f"Provide latest news and developments in the {business_type} industry"
        web_response = web_agent.run(web_prompt)
        agent_outputs.append(("Industry News", str(web_response)))
        progress_bar.progress(progress_step)

        # Tech Market Agent
        status_text.text("Analyzing market...")
        market_prompt = f"Based on the above news, provide detailed market analysis for {business_type} industry"
        market_response = tech_market_agent.run(market_prompt)
        agent_outputs.append(("Market Analysis", str(market_response)))
        progress_bar.progress(progress_step * 2)

        # Finance Agent
        status_text.text("Analyzing financials...")
        finance_prompt = f"Analyze financial metrics of key players in the {business_type} industry"
        finance_response = finance_agent.run(finance_prompt)
        agent_outputs.append(("Financial Analysis", str(finance_response)))
        progress_bar.progress(progress_step * 3)

        # Value Capture Agent
        status_text.text("Developing strategies...")
        value_prompt = f"Develop strategic recommendations for entering the {business_type} market"
        value_response = value_capture_agent.run(value_prompt)
        agent_outputs.append(("Strategic Recommendations", str(value_response)))
        progress_bar.progress(progress_step * 4)

        # Org Design Agent
        status_text.text("Designing organization...")
        org_prompt = f"Propose organizational structure for a {business_type} company"
        org_response = org_design_agent.run(org_prompt)
        agent_outputs.append(("Organizational Design", str(org_response)))
        progress_bar.progress(1.0)  # Changed from 100 to 1.0
        
        return agent_outputs
        
    except Exception as e:
        st.error(f"Error during analysis: {str(e)}")
        return None

def main():
    st.set_page_config(page_title="Business Analysis System", layout="wide")
    
    # Move input controls to sidebar
    with st.sidebar:
        st.title("Business Analysis System")
        st.write("Generate comprehensive business analysis reports using AI agents")
        
        business_type = st.text_input("What kind of company do you want to analyze?")
        generate_button = st.button("Generate Analysis")
    
    # Main area for results
    if generate_button:
        if not business_type:
            st.sidebar.warning("Please enter a business type")
            return
            
        progress_bar = st.sidebar.progress(0)
        status_text = st.sidebar.empty()
        
        agent_outputs = run_analysis(business_type, progress_bar, status_text)
        
        if agent_outputs:
            # Create tabs for different sections in main area
            tabs = st.tabs(["Report View", "Download Report"])
            
            with tabs[0]:
                for agent_name, output in agent_outputs:
                    with st.expander(agent_name, expanded=True):
                        st.markdown(output)
            
            with tabs[1]:
                # Generate Word document
                formatter = ReportFormatter()
                title = f"{business_type.title()} Industry Analysis Report"
                formatter.format_document(title, agent_outputs)
                
                # Get document as bytes
                docx_bytes = formatter.save_to_bytes()
                
                # Create download button
                st.download_button(
                    label="Download Full Report",
                    data=docx_bytes,
                    file_name=f"{business_type.lower().replace(' ', '_')}_analysis_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            progress_bar.empty()
            status_text.text("Analysis complete! You can now view the report above or download it.")

if __name__ == "__main__":
    main()
