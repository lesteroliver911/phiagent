from phi.agent import Agent
from phi.model.openai import OpenAIChat
from phi.tools.googlesearch import GoogleSearch
from phi.tools.yfinance import YFinanceTools
import inspect
import types
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import re
import os

# Define individual agents with clearer communication patterns
web_agent = Agent(
    name="Web Agent",
    role="Search the web for latest information and news",
    model=OpenAIChat(id="gpt-4o"),
    tools=[GoogleSearch()],
    instructions=[
        "Search for latest news and information about the given topic",
        "Provide 3-5 key findings with dates and sources",
        "Format information in clear, digestible bullet points"
    ],
    show_tool_calls=True,
    markdown=True,
)

finance_agent = Agent(
    name="Finance Agent",
    role="Analyze financial data and market trends",
    model=OpenAIChat(id="gpt-4o"),
    tools=[YFinanceTools(stock_price=True, analyst_recommendations=True, company_info=True)],
    instructions=[
        "Analyze financial metrics and market data",
        "Present data in clear tables",
        "Highlight key financial insights and trends"
    ],
    show_tool_calls=True,
    markdown=True,
)

tech_market_agent = Agent(
    name="Technology and Market Opportunity Expert",
    role="Analyze technology trends, market dynamics, and identify value creation opportunities",
    model=OpenAIChat(id="gpt-4o"),
    tools=[GoogleSearch()],
    instructions=[
        "Provide a structured market analysis with these specific sections:",
        
        "1. MARKET SIZE & GROWTH",
        "- Current global market size with specific dollar amount",
        "- Year-over-year growth rate (CAGR)",
        "- 5-year market size projection",
        "- Break down by major geographic regions",
        
        "2. MARKET SEGMENTS",
        "- List top 3-5 market segments with size/share",
        "- Identify fastest growing segments",
        "- Key drivers for each segment",
        
        "3. COMPETITIVE LANDSCAPE",
        "- Market share of top 5 players",
        "- Recent funding rounds and valuations",
        "- Key partnerships and acquisitions",
        
        "4. GROWTH DRIVERS & TRENDS",
        "- List specific technological advancements",
        "- Regulatory impacts",
        "- Customer demand patterns",
        
        "Important:",
        "- Always include specific numbers and dates",
        "- Cite sources for all data",
        "- Use bullet points for clarity",
        "- If exact figures unavailable, provide range estimates with justification"
    ],
    show_tool_calls=True,
    markdown=True,
)

value_capture_agent = Agent(
    name="Value Capture Strategist",
    role="Develop strategies for IP protection, market positioning, and competitive advantage",
    model=OpenAIChat(id="gpt-4o"),
    tools=[GoogleSearch()],
    instructions=[
        "Focus on IP protection strategies",
        "Develop market positioning recommendations",
        "Identify competitive advantages",
        "Provide actionable strategic recommendations",
        "Always include sources"
    ],
    show_tool_calls=True,
    markdown=True,
)

org_design_agent = Agent(
    name="Organizational Design Architect",
    role="Design optimal organizational structures and collaboration networks",
    model=OpenAIChat(id="gpt-4o"),
    tools=[GoogleSearch()],
    instructions=[
        "Design team structures and collaboration frameworks",
        "Optimize for innovation and value delivery",
        "Consider organizational culture and dynamics",
        "Provide practical implementation steps",
        "Always include sources"
    ],
    show_tool_calls=True,
    markdown=True,
)

# Combine agents with clear workflow
agent_team = Agent(
    name="Product Management Research Team",
    team=[
        web_agent, 
        finance_agent, 
        tech_market_agent, 
        value_capture_agent, 
        org_design_agent
    ],
    instructions=[
        "Follow this exact sequence:",
        "1. Web Agent: Gather latest news and developments",
        "2. Tech Market Agent: Using news context, provide detailed market analysis",
        "3. Finance Agent: Add financial metrics of key players",
        "4. Value Capture Agent: Develop strategies based on market analysis",
        "5. Org Design Agent: Propose structure based on market size and opportunity",
        
        "Ensure each agent builds upon previous agents' findings",
        "All market size claims must include sources",
        "Present data in tables where appropriate",
        "Highlight key opportunities and risks"
    ],
    show_tool_calls=True,
    markdown=True,
)

# Function signature inspection utility
def inspect_function(func):
    if callable(func):
        if isinstance(func, types.BuiltinFunctionType):
            print("Cannot inspect signature of built-in type.")
        else:
            try:
                argspec = inspect.signature(func)
                print(f"Signature: {argspec}")
            except ValueError:
                print("Signature could not be retrieved.")
    else:
        raise ValueError("Provided object is not callable.")

class ReportFormatter:
    def __init__(self):
        self.document = Document()
        # Set up document styling
        style = self.document.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
    def clean_output(self, text):
        # Remove the "Running:" sections and debug info
        cleaned = re.sub(r'Running:.*?(?=\n\n)', '', text, flags=re.DOTALL)
        cleaned = re.sub(r'• transfer_task_to.*?\.\.\..*?\n', '', cleaned)
        # Remove any empty lines at the start
        cleaned = re.sub(r'^\s*\n', '', cleaned)
        return cleaned.strip()
        
    def add_section(self, title, content):
        """Add a new section with heading and content"""
        # Add section heading
        section_heading = self.document.add_heading(level=1)
        heading_run = section_heading.add_run(title)
        heading_run.bold = True
        
        # Add content
        cleaned_content = self.clean_output(content)
        self.document.add_paragraph(cleaned_content)
        
        # Add separator
        self.document.add_paragraph('_' * 50)

    def format_document(self, title, agent_outputs):
        """Format document with multiple agent outputs"""
        # Add main title
        title_heading = self.document.add_heading(level=0)
        title_run = title_heading.add_run(title)
        title_run.bold = True
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_paragraph = self.document.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_paragraph.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        
        # Add separator
        self.document.add_paragraph('_' * 50)
        
        # Add each agent's output as a separate section
        for agent_name, output in agent_outputs:
            self.add_section(f"{agent_name} Analysis", output)

    def save(self, business_type):
        # Create reports directory if it doesn't exist
        if not os.path.exists('reports'):
            os.makedirs('reports')
            
        # Format filename
        formatted_business = business_type.lower().replace(' ', '_')
        filename = f"reports/{formatted_business}_market_analysis_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        
        self.document.save(filename)
        return filename

# Function to ask the user for input and run the agent team
def start_business_analysis():
    business_type = input("What kind of company do you want to create? ")
    
    # Create base prompts for each agent
    web_prompt = f"Provide latest news and developments in the {business_type} industry"
    market_prompt = f"Based on the above news, provide detailed market analysis for {business_type} industry"
    finance_prompt = f"Analyze financial metrics of key players in the {business_type} industry"
    value_prompt = f"Develop strategic recommendations for entering the {business_type} market"
    org_prompt = f"Propose organizational structure for a {business_type} company"
    
    try:
        print(f"\n=== {business_type.title()} Industry Analysis ===\n")
        
        # Collect outputs from each agent
        agent_outputs = []
        
        # Web Agent
        print("\nGathering latest news...")
        web_response = web_agent.run(web_prompt)
        agent_outputs.append(("Industry News", str(web_response)))
        
        # Tech Market Agent
        print("\nAnalyzing market...")
        market_response = tech_market_agent.run(market_prompt)
        agent_outputs.append(("Market Analysis", str(market_response)))
        
        # Finance Agent
        print("\nAnalyzing financials...")
        finance_response = finance_agent.run(finance_prompt)
        agent_outputs.append(("Financial Analysis", str(finance_response)))
        
        # Value Capture Agent
        print("\nDeveloping strategies...")
        value_response = value_capture_agent.run(value_prompt)
        agent_outputs.append(("Strategic Recommendations", str(value_response)))
        
        # Org Design Agent
        print("\nDesigning organization...")
        org_response = org_design_agent.run(org_prompt)
        agent_outputs.append(("Organizational Design", str(org_response)))
        
        # Create and save the document
        formatter = ReportFormatter()
        title = f"{business_type.title()} Industry Analysis Report"
        formatter.format_document(title, agent_outputs)
        filename = formatter.save(business_type)
        
        print(f"\nDetailed report saved as: {filename}")
        # Open the file automatically
        os.system(f"start {filename}" if os.name == 'nt' else f"open {filename}")
        
    except Exception as e:
        print(f"\nError during analysis: {str(e)}")
        print("Try narrowing the industry scope or checking the input format.")

# Start the process
start_business_analysis()

def test_web_agent():
    test_prompt = "Search for the top 3 latest news items about AI startups."
    print("\n=== Web Agent Search Results ===\n")
    
    try:
        response = web_agent.print_response(test_prompt, stream=True)
        
        # Format and save to Word document
        formatter = ReportFormatter()
        formatter.format_document(str(response), "AI Startups News Report")
        filename = f"ai_startups_news_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        formatter.save(filename)
        print(f"\nReport saved as: {filename}")
        
    except Exception as e:
        print(f"\nError in web agent test: {str(e)}")

# Call the test function
test_web_agent()